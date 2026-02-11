# train_fuzzy_search.py

import os
import re
import pickle
import numpy as np
import pandas as pd
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Union, Tuple, Set, Optional

# NLP and vectorization tools
import nltk
from nltk.tokenize import word_tokenize
from nltk.stem import PorterStemmer, WordNetLemmatizer
from nltk.corpus import stopwords
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

# PDF, DOCX, and other document handling
import fitz  # PyMuPDF
import docx
from bs4 import BeautifulSoup
import csv
import xlrd

# Fuzzy matching
from fuzzywuzzy import fuzz, process

# Machine learning
from sklearn.decomposition import TruncatedSVD
from sklearn.ensemble import RandomForestClassifier
from sklearn.model_selection import train_test_split
from sklearn.metrics import accuracy_score, classification_report
from sklearn.pipeline import Pipeline
from sklearn.base import BaseEstimator, TransformerMixin

# Streamlit for the app interface
import streamlit as st

# Download necessary NLTK resources
def download_nltk_resources():
    """Download required NLTK resources for text processing."""
    resources = ['punkt', 'wordnet', 'stopwords']
    for resource in resources:
        try:
            nltk.data.find(f'tokenizers/{resource}')
        except LookupError:
            nltk.download(resource, quiet=True)

# Document Processing Classes
class DocumentProcessor:
    """Base class for document processing."""
    
    def __init__(self):
        self.supported_extensions = []
    
    def can_process(self, file_path: str) -> bool:
        """Check if this processor can handle the file type."""
        ext = os.path.splitext(file_path)[1].lower()
        return ext in self.supported_extensions
    
    def extract_text(self, file_path: str) -> str:
        """Extract text from document."""
        raise NotImplementedError("Subclasses must implement extract_text")
    
    def extract_metadata(self, file_path: str) -> Dict:
        """Extract metadata from document."""
        return {
            'filename': os.path.basename(file_path),
            'path': file_path,
            'size': os.path.getsize(file_path),
            'last_modified': datetime.fromtimestamp(os.path.getmtime(file_path)).isoformat(),
            'extension': os.path.splitext(file_path)[1].lower()
        }

class PDFProcessor(DocumentProcessor):
    """Process PDF documents."""
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = ['.pdf']
    
    def extract_text(self, file_path: str) -> str:
        """Extract text from PDF."""
        text = ""
        try:
            with fitz.open(file_path) as doc:
                for page in doc:
                    text += page.get_text()
        except Exception as e:
            print(f"Error processing PDF {file_path}: {e}")
        return text
    
    def extract_metadata(self, file_path: str) -> Dict:
        """Extract metadata from PDF."""
        metadata = super().extract_metadata(file_path)
        try:
            with fitz.open(file_path) as doc:
                metadata.update({
                    'page_count': len(doc),
                    'title': doc.metadata.get('title', ''),
                    'author': doc.metadata.get('author', ''),
                    'subject': doc.metadata.get('subject', ''),
                    'keywords': doc.metadata.get('keywords', ''),
                    'creator': doc.metadata.get('creator', ''),
                    'producer': doc.metadata.get('producer', ''),
                })
        except Exception as e:
            print(f"Error extracting PDF metadata from {file_path}: {e}")
        return metadata

class DocxProcessor(DocumentProcessor):
    """Process DOCX documents."""
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = ['.docx']
    
    def extract_text(self, file_path: str) -> str:
        """Extract text from DOCX."""
        text = ""
        try:
            doc = docx.Document(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
        except Exception as e:
            print(f"Error processing DOCX {file_path}: {e}")
        return text
    
    def extract_metadata(self, file_path: str) -> Dict:
        """Extract metadata from DOCX."""
        metadata = super().extract_metadata(file_path)
        try:
            doc = docx.Document(file_path)
            core_props = doc.core_properties
            metadata.update({
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'keywords': core_props.keywords or '',
                'created': core_props.created.isoformat() if core_props.created else '',
                'modified': core_props.modified.isoformat() if core_props.modified else '',
                'category': core_props.category or '',
                'comments': core_props.comments or '',
            })
        except Exception as e:
            print(f"Error extracting DOCX metadata from {file_path}: {e}")
        return metadata

class TxtProcessor(DocumentProcessor):
    """Process TXT documents."""
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = ['.txt']
    
    def extract_text(self, file_path: str) -> str:
        """Extract text from TXT."""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                return f.read()
        except Exception as e:
            print(f"Error processing TXT {file_path}: {e}")
            return ""

class HtmlProcessor(DocumentProcessor):
    """Process HTML documents."""
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = ['.html', '.htm']
    
    def extract_text(self, file_path: str) -> str:
        """Extract text from HTML."""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                soup = BeautifulSoup(f.read(), 'html.parser')
                # Remove script and style elements
                for script in soup(["script", "style"]):
                    script.decompose()
                # Get text
                text = soup.get_text(separator=' ', strip=True)
                # Break into lines and remove leading and trailing space
                lines = (line.strip() for line in text.splitlines())
                # Break multi-headlines into a line each
                chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                # Drop blank lines
                text = '\n'.join(chunk for chunk in chunks if chunk)
                return text
        except Exception as e:
            print(f"Error processing HTML {file_path}: {e}")
            return ""

class CsvProcessor(DocumentProcessor):
    """Process CSV documents."""
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = ['.csv']
    
    def extract_text(self, file_path: str) -> str:
        """Extract text from CSV."""
        text = ""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                reader = csv.reader(f)
                for row in reader:
                    text += " ".join(str(cell) for cell in row) + "\n"
        except Exception as e:
            print(f"Error processing CSV {file_path}: {e}")
        return text

class XlsProcessor(DocumentProcessor):
    """Process Excel documents."""
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = ['.xls', '.xlsx']
    
    def extract_text(self, file_path: str) -> str:
        """Extract text from Excel."""
        text = ""
        try:
            workbook = xlrd.open_workbook(file_path)
            for sheet in workbook.sheets():
                for row in range(sheet.nrows):
                    values = []
                    for col in range(sheet.ncols):
                        values.append(str(sheet.cell(row, col).value))
                    text += " ".join(values) + "\n"
        except Exception as e:
            print(f"Error processing Excel {file_path}: {e}")
        return text

# Text Preprocessing
class TextPreprocessor:
    """Clean and preprocess text for better search results."""
    
    def __init__(self):
        download_nltk_resources()
        self.stop_words = set(stopwords.words('english'))
        self.stemmer = PorterStemmer()
        self.lemmatizer = WordNetLemmatizer()
    
    def clean_text(self, text: str) -> str:
        """Basic text cleaning."""
        if not text:
            return ""
        # Convert to lowercase
        text = text.lower()
        # Remove special characters and digits
        text = re.sub(r'[^\w\s]', ' ', text)
        text = re.sub(r'\d+', ' ', text)
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text).strip()
        return text
    
    def tokenize(self, text: str) -> List[str]:
        """Tokenize text into words."""
        return word_tokenize(text)
    
    def remove_stopwords(self, tokens: List[str]) -> List[str]:
        """Remove common stopwords."""
        return [token for token in tokens if token not in self.stop_words]
    
    def stem_tokens(self, tokens: List[str]) -> List[str]:
        """Apply stemming to tokens."""
        return [self.stemmer.stem(token) for token in tokens]
    
    def lemmatize_tokens(self, tokens: List[str]) -> List[str]:
        """Apply lemmatization to tokens."""
        return [self.lemmatizer.lemmatize(token) for token in tokens]
    
    def preprocess(self, text: str, use_stemming: bool = False) -> str:
        """Full preprocessing pipeline."""
        if not text:
            return ""
        # Clean text
        cleaned_text = self.clean_text(text)
        # Tokenize
        tokens = self.tokenize(cleaned_text)
        # Remove stopwords
        tokens = self.remove_stopwords(tokens)
        # Apply stemming or lemmatization
        if use_stemming:
            tokens = self.stem_tokens(tokens)
        else:
            tokens = self.lemmatize_tokens(tokens)
        # Join tokens back to text
        processed_text = ' '.join(tokens)
        return processed_text

# Document Indexing
class DocumentIndexer:
    """Index documents for fast retrieval."""
    
    def __init__(self, preprocessor: TextPreprocessor = None):
        self.preprocessor = preprocessor or TextPreprocessor()
        self.document_processors = [
            PDFProcessor(),
            DocxProcessor(),
            TxtProcessor(),
            HtmlProcessor(),
            CsvProcessor(),
            XlsProcessor(),
        ]
        self.documents = []
        self.vectorizer = TfidfVectorizer(
            max_df=0.9,
            min_df=2,
            max_features=10000,
            ngram_range=(1, 2)
        )
        self.svd = TruncatedSVD(n_components=100)
        self.document_vectors = None
        self.reduced_vectors = None
        self.index_created = False
    
    def get_processor_for_file(self, file_path: str) -> DocumentProcessor:
        """Get the appropriate processor for a file."""
        for processor in self.document_processors:
            if processor.can_process(file_path):
                return processor
        return None
    
    def index_document(self, file_path: str) -> Dict:
        """Process and index a single document."""
        processor = self.get_processor_for_file(file_path)
        if not processor:
            print(f"No processor found for {file_path}")
            return None
        
        text = processor.extract_text(file_path)
        metadata = processor.extract_metadata(file_path)
        processed_text = self.preprocessor.preprocess(text)
        
        document = {
            'id': len(self.documents),
            'file_path': file_path,
            'metadata': metadata,
            'raw_text': text,
            'processed_text': processed_text,
            'chunks': self._chunk_text(text, chunk_size=1000, overlap=200),
        }
        
        self.documents.append(document)
        self.index_created = False
        return document
    
    def _chunk_text(self, text: str, chunk_size: int = 1000, overlap: int = 200) -> List[Dict]:
        """Split text into overlapping chunks for better search."""
        chunks = []
        if not text:
            return chunks
        
        words = text.split()
        i = 0
        while i < len(words):
            chunk_words = words[i:i + chunk_size]
            chunk_text = ' '.join(chunk_words)
            chunks.append({
                'text': chunk_text,
                'processed_text': self.preprocessor.preprocess(chunk_text),
                'position': i,
            })
            i += chunk_size - overlap
        
        return chunks
    
    def index_directory(self, directory_path: str, recursive: bool = True) -> List[Dict]:
        """Process and index all documents in a directory."""
        indexed = []
        for root, _, files in os.walk(directory_path):
            for file in files:
                file_path = os.path.join(root, file)
                document = self.index_document(file_path)
                if document:
                    indexed.append(document)
            if not recursive:
                break
        return indexed
    
    def build_vector_index(self):
        """Build vector representations of all documents."""
        if not self.documents:
            return
        
        # Extract all processed texts
        texts = [doc['processed_text'] for doc in self.documents]
        
        # Additional texts from chunks for more granular search
        for doc in self.documents:
            for chunk in doc['chunks']:
                texts.append(chunk['processed_text'])
        
        # Fit and transform
        self.document_vectors = self.vectorizer.fit_transform(texts)
        
        # Apply dimensionality reduction
        self.reduced_vectors = self.svd.fit_transform(self.document_vectors)
        
        self.index_created = True
    
    def save_index(self, path: str):
        """Save the document index to disk."""
        if not self.index_created:
            self.build_vector_index()
        
        data = {
            'documents': self.documents,
            'vectorizer': self.vectorizer,
            'svd': self.svd,
            'document_vectors': self.document_vectors,
            'reduced_vectors': self.reduced_vectors,
        }
        
        with open(path, 'wb') as f:
            pickle.dump(data, f)
    
    def load_index(self, path: str):
        """Load the document index from disk."""
        with open(path, 'rb') as f:
            data = pickle.load(f)
        
        self.documents = data['documents']
        self.vectorizer = data['vectorizer']
        self.svd = data['svd']
        self.document_vectors = data['document_vectors']
        self.reduced_vectors = data['reduced_vectors']
        self.index_created = True

# Search Engine
class FuzzySearchEngine:
    """Search engine with fuzzy matching capabilities."""
    
    def __init__(self, indexer: DocumentIndexer = None):
        self.indexer = indexer or DocumentIndexer()
        self.preprocessor = self.indexer.preprocessor
        self.similarity_threshold = 0.3
        self.fuzzy_ratio_threshold = 70
    
    def search(self, query: str, top_k: int = 10) -> List[Dict]:
        """
        Search for documents matching the query.
        
        Args:
            query: Search query
            top_k: Number of results to return
            
        Returns:
            List of documents with relevance scores
        """
        if not self.indexer.index_created:
            self.indexer.build_vector_index()
        
        # Preprocess query
        processed_query = self.preprocessor.preprocess(query)
        
        # Vector search
        query_vector = self.indexer.vectorizer.transform([processed_query])
        query_reduced = self.indexer.svd.transform(query_vector)
        
        # Calculate similarities
        similarities = cosine_similarity(query_reduced, self.indexer.reduced_vectors)[0]
        
        # Combine with fuzzy search for better results
        fuzzy_scores = self._fuzzy_search(query, processed_query)
        
        # Create results
        results = []
        doc_count = len(self.indexer.documents)
        
        for i, sim in enumerate(similarities):
            # Check if this is a document or chunk
            if i < doc_count:
                # This is a document
                doc_idx = i
                is_chunk = False
                chunk_idx = None
            else:
                # This is a chunk
                chunk_count = 0
                for d, doc in enumerate(self.indexer.documents):
                    new_chunk_count = chunk_count + len(doc['chunks'])
                    if i - doc_count < new_chunk_count:
                        doc_idx = d
                        is_chunk = True
                        chunk_idx = i - doc_count - chunk_count
                        break
                    chunk_count = new_chunk_count
            
            # Skip if similarity is too low
            if sim < self.similarity_threshold:
                continue
            
            doc = self.indexer.documents[doc_idx]
            
            # Get fuzzy score
            fuzzy_score = fuzzy_scores.get(doc['id'], 0)
            
            # Calculate combined score
            combined_score = 0.7 * sim + 0.3 * (fuzzy_score / 100)
            
            # Skip if combined score is too low
            if combined_score < self.similarity_threshold:
                continue
            
            # Create result
            result = {
                'document': doc,
                'vector_similarity': float(sim),
                'fuzzy_score': fuzzy_score,
                'combined_score': combined_score,
                'is_chunk': is_chunk,
            }
            
            # Add chunk info if this is a chunk
            if is_chunk:
                result['chunk'] = doc['chunks'][chunk_idx]
            
            results.append(result)
        
        # Sort by combined score and take top_k
        results.sort(key=lambda x: x['combined_score'], reverse=True)
        return results[:top_k]
    
    def _fuzzy_search(self, raw_query: str, processed_query: str) -> Dict[int, int]:
        """
        Perform fuzzy search on documents.
        
        Returns:
            Dictionary mapping document ID to fuzzy match score
        """
        scores = {}
        for doc in self.indexer.documents:
            # Check raw text
            ratio1 = fuzz.partial_ratio(raw_query, doc['raw_text'])
            # Check processed text
            ratio2 = fuzz.token_sort_ratio(processed_query, doc['processed_text'])
            # Take the max score
            scores[doc['id']] = max(ratio1, ratio2)
        return scores
    
    def search_by_metadata(self, metadata_filters: Dict, top_k: int = 10) -> List[Dict]:
        """
        Search for documents by metadata.
        
        Args:
            metadata_filters: Dictionary of metadata field-value pairs
            top_k: Number of results to return
            
        Returns:
            List of documents matching the metadata filters
        """
        results = []
        
        for doc in self.indexer.documents:
            match_score = 0
            for field, value in metadata_filters.items():
                if field in doc['metadata']:
                    if isinstance(value, str) and isinstance(doc['metadata'][field], str):
                        # Fuzzy match for strings
                        ratio = fuzz.partial_ratio(value.lower(), doc['metadata'][field].lower())
                        if ratio >= self.fuzzy_ratio_threshold:
                            match_score += ratio / 100.0
                    elif doc['metadata'][field] == value:
                        # Exact match for other types
                        match_score += 1
            
            if match_score > 0:
                results.append({
                    'document': doc,
                    'metadata_match_score': match_score,
                    'is_chunk': False,
                })
        
        # Sort by match score and take top_k
        results.sort(key=lambda x: x['metadata_match_score'], reverse=True)
        return results[:top_k]

# Custom ML models for document classification
class DocumentClassifier:
    """Classify documents based on their content."""
    
    def __init__(self, preprocessor: TextPreprocessor = None):
        self.preprocessor = preprocessor or TextPreprocessor()
        self.vectorizer = TfidfVectorizer(
            max_df=0.9,
            min_df=2,
            max_features=5000,
            ngram_range=(1, 2)
        )
        self.clf = RandomForestClassifier(n_estimators=100, random_state=42)
        self.pipeline = Pipeline([
            ('vectorizer', self.vectorizer),
            ('classifier', self.clf)
        ])
        self.classes = []
        self.trained = False
    
    def train(self, documents: List[Dict], labels: List[str]):
        """
        Train the classifier on documents.
        
        Args:
            documents: List of document dictionaries
            labels: Document labels/categories
        """
        if not documents or not labels:
            return
        
        # Extract processed texts
        texts = [doc['processed_text'] for doc in documents]
        
        # Get unique classes
        self.classes = sorted(set(labels))
        
        # Train-test split
        X_train, X_test, y_train, y_test = train_test_split(
            texts, labels, test_size=0.2, random_state=42
        )
        
        # Train the pipeline
        self.pipeline.fit(X_train, y_train)
        
        # Evaluate
        y_pred = self.pipeline.predict(X_test)
        accuracy = accuracy_score(y_test, y_pred)
        report = classification_report(y_test, y_pred, target_names=self.classes)
        
        print(f"Classifier accuracy: {accuracy:.4f}")
        print(report)
        
        self.trained = True
    
    def predict(self, document: Dict) -> Tuple[str, float]:
        """
        Predict the category of a document.
        
        Args:
            document: Document dictionary
            
        Returns:
            Tuple of (predicted_class, confidence)
        """
        if not self.trained:
            return None, 0.0
        
        # Extract processed text
        text = document['processed_text']
        
        # Predict
        proba = self.pipeline.predict_proba([text])[0]
        class_idx = proba.argmax()
        confidence = proba[class_idx]
        
        return self.classes[class_idx], confidence
    
    def save_model(self, path: str):
        """Save the trained model to disk."""
        if not self.trained:
            print("Model not trained, nothing to save.")
            return
        
        data = {
            'vectorizer': self.vectorizer,
            'classifier': self.clf,
            'classes': self.classes,
            'trained': self.trained,
        }
        
        with open(path, 'wb') as f:
            pickle.dump(data, f)
    
    def load_model(self, path: str):
        """Load the trained model from disk."""
        with open(path, 'rb') as f:
            data = pickle.load(f)
        
        self.vectorizer = data['vectorizer']
        self.clf = data['classifier']
        self.classes = data['classes']
        self.trained = data['trained']
        
        # Recreate pipeline
        self.pipeline = Pipeline([
            ('vectorizer', self.vectorizer),
            ('classifier', self.clf)
        ])

# Streamlit App
def create_streamlit_app():
    """Create the Streamlit app."""
    st.set_page_config(
        page_title="AI-Powered Document Search",
        page_icon="📄",
        layout="wide"
    )
    
    # Initialize session state
    if 'indexer' not in st.session_state:
        st.session_state.indexer = DocumentIndexer()
    if 'search_engine' not in st.session_state:
        st.session_state.search_engine = FuzzySearchEngine(st.session_state.indexer)
    if 'classifier' not in st.session_state:
        st.session_state.classifier = DocumentClassifier()
    if 'documents' not in st.session_state:
        st.session_state.documents = []
    if 'index_path' not in st.session_state:
        st.session_state.index_path = "document_index.pkl"
    if 'model_path' not in st.session_state:
        st.session_state.model_path = "document_classifier.pkl"
    
    # App title
    st.title("AI-Powered Document Search")
    
    # Sidebar for settings and controls
    with st.sidebar:
        st.header("Settings")
        
        # Mode selection
        mode = st.radio("Mode", ["Search", "Index", "Train", "Settings"])
        
        # Index settings
        if mode == "Index" or mode == "Settings":
            st.subheader("Index Settings")
            index_path = st.text_input("Index Path", value=st.session_state.index_path)
            if index_path != st.session_state.index_path:
                st.session_state.index_path = index_path
            
            # Load index button
            if st.button("Load Index", key="load_index"):
                try:
                    with st.spinner("Loading index..."):
                        st.session_state.indexer.load_index(st.session_state.index_path)
                        st.session_state.documents = st.session_state.indexer.documents
                    st.success(f"Loaded {len(st.session_state.documents)} documents from index.")
                except Exception as e:
                    st.error(f"Error loading index: {e}")
            
            # Save index button
            if st.button("Save Index", key="save_index"):
                try:
                    with st.spinner("Saving index..."):
                        st.session_state.indexer.save_index(st.session_state.index_path)
                    st.success("Index saved successfully.")
                except Exception as e:
                    st.error(f"Error saving index: {e}")
        
        # Classifier settings
        if mode == "Train" or mode == "Settings":
            st.subheader("Classifier Settings")
            model_path = st.text_input("Model Path", value=st.session_state.model_path)
            if model_path != st.session_state.model_path:
                st.session_state.model_path = model_path
            
            # Load model button
            if st.button("Load Model", key="load_model"):
                try:
                    with st.spinner("Loading model..."):
                        st.session_state.classifier.load_model(st.session_state.model_path)
                    st.success("Model loaded successfully.")
                except Exception as e:
                    st.error(f"Error loading model: {e}")
            
            # Save model button
            if st.button("Save Model", key="save_model"):
                try:
                    with st.spinner("Saving model..."):
                        st.session_state.classifier.save_model(st.session_state.model_path)
                    st.success("Model saved successfully.")
                except Exception as e:
                    st.error(f"Error saving model: {e}")
        
        # About
        st.markdown("---")
        st.markdown("### About")
        st.markdown(
            "This app uses AI to search and classify documents. "
            "It works completely offline and can handle various document types."
        )
    
    # Main content based on selected mode
    if mode == "Search":
        create_search_ui()
    elif mode == "Index":
        create_index_ui()
    elif mode == "Train":
        create_train_ui()
    elif mode == "Settings":
        create_settings_ui()

def create_search_ui():
    """Create the search UI."""
    st.header("Document Search")
    
    # Search form
    with st.form("search_form"):
        query = st.text_input("Search Query", placeholder="Enter your search terms...")
        col1, col2 = st.columns(2)
        
        with col1:
            use_fuzzy = st.checkbox("Use Fuzzy Matching", value=True)
            num_results = st.slider("Number of Results", min_value=1, max_value=50, value=10)
        
        with col2:
            # Metadata filters
            st.markdown("Metadata Filters (Optional)")
            filter_filename = st.text_input("Filename Contains", "")
            filter_extension = st.text_input("File Extension", "")
        
        # Search button
        search_button = st.form_submit_button("Search")
    
    # Process search
    if search_button and query:
        # Check if index exists
        if not st.session_state.indexer.documents:
            st.warning("No documents indexed. Please go to the Index tab to add documents.")
            return
        
        # Create metadata filters
        metadata_filters = {}
        if filter_filename:
            metadata_filters['filename'] = filter_filename
        if filter_extension:
            metadata_filters['extension'] = filter_extension
        
        # Perform search
        with st.spinner("Searching..."):
            # Vector search
            results = st.session_state.search_engine.search(query, top_k=num_results)
            
            # Metadata search if filters provided
            if metadata_filters:
                metadata_results = st.session_state.search_engine.search_by_metadata(
                    metadata_filters, top_k=num_results
                )
                
                # Combine results
                all_ids = set(r['document']['id'] for r in results)
                for r in metadata_results:
                    if r['document']['id'] not in all_ids:
                        results.append(r)
                
                # Sort again
                results.sort(key=lambda x: x.get('combined_score', 0) + x.get('metadata_match_score', 0), reverse=True)
                results = results[:num_results]
        
        # Display results
        if results:
            st.subheader(f"Found {len(results)} results")
            
            for i, result in enumerate(results):
                doc = result['document']
                score = result.get('combined_score', 0) + result.get('metadata_match_score', 0)
                
                with st.expander(
                    f"{i+1}. {doc['metadata']['filename']} (Score: {score:.2f})"
                ):
                    # Display metadata
                    st.markdown("**Metadata:**")
                    metadata_df = pd.DataFrame([doc['metadata']])
                    st.dataframe(metadata_df)
                    
                    # Display snippet
                    st.markdown("**Content Snippet:**")
                    if result.get('is_chunk', False):
                        snippet = result['chunk']['text']
                        position = result['chunk']['position']
                        st.markdown(f"Position: {position}")
                    else:
                        # Extract a snippet around the matching term
                        text = doc['raw_text']
                        query_terms = query.lower().split()
                        snippet = ""
                        for term in query_terms:
                            pos = text.lower().find(term)
                            if pos >= 0:
                                start = max(0, pos - 100)
                                end = min(len(text), pos + len(term) + 100)
                                term_snippet = text[start:end]
                                snippet += f"...{term_snippet}...\n\n"
                        if not snippet:
                            # If no term found, just take the first 300 chars
                            snippet = text[:300] + "..."
                    
                    st.text_area("", value=snippet, height=150)
                    
                    # Open document button
                    st.markdown(f"[Open Document]({doc['file_path']})")
                    
                    # If classifier is trained, show classification
                    if st.session_state.classifier.trained:
                        category, confidence = st.session_state.classifier.predict(doc)
                        st.markdown(f"**Category:** {category} (Confidence: {confidence:.2f})")
        else:
            st.info("No results found. Try a different query or adjust your search settings.")

def create_index_ui():
    """Create the indexing UI."""
    st.header("Document Indexing")
    
    # Document statistics
    if st.session_state.documents:
        st.subheader("Index Statistics")
        
        # Count document types
        extensions = {}
        for doc in st.session_state.documents:
            ext = doc['metadata']['extension']
            extensions[ext] = extensions.get(ext, 0) + 1
        
        # Create chart data
        chart_data = pd.DataFrame({
            'Extension': list(extensions.keys()),
            'Count': list(extensions.values())
        })
        
        # Display statistics
        col1, col2 = st.columns(2)
        with col1:
            st.metric("Total Documents", len(st.session_state.documents))
            st.metric("Total File Types", len(extensions))
        
        with col2:
            st.bar_chart(chart_data.set_index('Extension'))
    
    # Index directory form
    with st.form("index_directory_form"):
        st.subheader("Index Directory")
        directory_path = st.text_input("Directory Path", placeholder="Enter path to directory containing documents")
        recursive = st.checkbox("Include Subdirectories", value=True)
        
        # Submit button
        submit_button = st.form_submit_button("Index Directory")
    
    # Process form
    if submit_button and directory_path:
        # Validate directory
        if not os.path.isdir(directory_path):
            st.error(f"Directory '{directory_path}' not found.")
            return
        
        # Index directory
        with st.spinner(f"Indexing directory '{directory_path}'..."):
            try:
                indexed = st.session_state.indexer.index_directory(directory_path, recursive=recursive)
                st.session_state.documents = st.session_state.indexer.documents
                st.success(f"Indexed {len(indexed)} documents.")
                
                # Build vector index
                st.session_state.indexer.build_vector_index()
                st.success("Built vector index.")
                
                # Suggest saving
                st.info("Don't forget to save your index in the sidebar.")
            except Exception as e:
                st.error(f"Error indexing directory: {str(e)}")
    
    # Index single file form
    with st.form("index_file_form"):
        st.subheader("Index Single File")
        file_path = st.text_input("File Path", placeholder="Enter path to file")
        
        # Submit button
        submit_button = st.form_submit_button("Index File")
    
    # Process form
    if submit_button and file_path:
        # Validate file
        if not os.path.isfile(file_path):
            st.error(f"File '{file_path}' not found.")
            return
        
        # Index file
        with st.spinner(f"Indexing file '{file_path}'..."):
            try:
                document = st.session_state.indexer.index_document(file_path)
                if document:
                    st.session_state.documents = st.session_state.indexer.documents
                    st.success(f"Indexed file '{file_path}'.")
                    
                    # Build vector index
                    st.session_state.indexer.build_vector_index()
                    st.success("Built vector index.")
                    
                    # Suggest saving
                    st.info("Don't forget to save your index in the sidebar.")
                else:
                    st.warning(f"Could not index file '{file_path}'. Unsupported file type.")
            except Exception as e:
                st.error(f"Error indexing file: {str(e)}")
    
    # List indexed documents
    if st.session_state.documents:
        st.subheader("Indexed Documents")
        
        # Create dataframe
        data = [{
            'ID': doc['id'],
            'Filename': doc['metadata']['filename'],
            'Path': doc['metadata']['path'],
            'Size (KB)': round(doc['metadata']['size'] / 1024, 2),
            'Type': doc['metadata']['extension'],
        } for doc in st.session_state.documents]
        
        df = pd.DataFrame(data)
        st.dataframe(df)

def create_train_ui():
    """Create the training UI."""
    st.header("Model Training")
    
    # Check if we have documents
    if not st.session_state.documents:
        st.warning("No documents indexed. Please go to the Index tab to add documents first.")
        return
    
    # Training data
    st.subheader("Create Training Data")
    st.markdown(
        "Assign categories to documents to train the classifier. "
        "You can manually categorize documents or use rules to assign categories automatically."
    )
    
    # Manual categorization
    with st.expander("Manual Categorization"):
        # Create a selection for document
        document_options = [f"{doc['id']}: {doc['metadata']['filename']}" for doc in st.session_state.documents]
        selected_doc = st.selectbox("Select Document", document_options)
        
        if selected_doc:
            # Get document ID
            doc_id = int(selected_doc.split(":")[0])
            doc = next((d for d in st.session_state.documents if d['id'] == doc_id), None)
            
            if doc:
                # Display document info
                st.markdown(f"**Filename:** {doc['metadata']['filename']}")
                st.markdown(f"**Path:** {doc['metadata']['path']}")
                st.markdown(f"**Size:** {round(doc['metadata']['size'] / 1024, 2)} KB")
                
                # Display snippet
                st.markdown("**Content Snippet:**")
                st.text_area("", value=doc['raw_text'][:500] + "...", height=200)
                
                # Category input
                category = st.text_input("Category", key=f"category_{doc_id}")
                
                # Save button
                if st.button("Assign Category", key=f"assign_{doc_id}"):
                    if category:
                        # Add category to document
                        doc['category'] = category
                        st.success(f"Assigned category '{category}' to document.")
                    else:
                        st.error("Please enter a category.")
    
    # Automatic categorization
    with st.expander("Automatic Categorization"):
        st.markdown("Create rules to automatically assign categories to documents.")
        
        # Rule types
        rule_type = st.selectbox(
            "Rule Type",
            ["Filename Contains", "Extension Is", "Content Contains"]
        )
        
        # Rule value
        rule_value = st.text_input("Rule Value")
        
        # Category
        auto_category = st.text_input("Category to Assign")
        
        # Apply button
        if st.button("Apply Rule"):
            if not rule_value or not auto_category:
                st.error("Please enter both a rule value and category.")
            else:
                # Apply rule to documents
                count = 0
                for doc in st.session_state.documents:
                    if rule_type == "Filename Contains" and rule_value.lower() in doc['metadata']['filename'].lower():
                        doc['category'] = auto_category
                        count += 1
                    elif rule_type == "Extension Is" and rule_value.lower() == doc['metadata']['extension'].lower():
                        doc['category'] = auto_category
                        count += 1
                    elif rule_type == "Content Contains" and rule_value.lower() in doc['raw_text'].lower():
                        doc['category'] = auto_category
                        count += 1
                
                st.success(f"Assigned category '{auto_category}' to {count} documents.")
    
    # Train model
    st.subheader("Train Classifier")
    
    # Check if we have categorized documents
    categorized_docs = [doc for doc in st.session_state.documents if 'category' in doc]
    
    if not categorized_docs:
        st.warning("No categorized documents. Please assign categories to documents first.")
    else:
        st.markdown(f"Found {len(categorized_docs)} categorized documents.")
        
        # Display categories
        categories = {}
        for doc in categorized_docs:
            categories[doc['category']] = categories.get(doc['category'], 0) + 1
        
        # Create chart data
        chart_data = pd.DataFrame({
            'Category': list(categories.keys()),
            'Count': list(categories.values())
        })
        
        # Display chart
        st.bar_chart(chart_data.set_index('Category'))
        
        # Training parameters
        st.markdown("### Training Parameters")
        test_size = st.slider("Test Size", min_value=0.1, max_value=0.5, value=0.2, step=0.05)
        
        # Train button
        if st.button("Train Classifier"):
            with st.spinner("Training classifier..."):
                try:
                    # Prepare data
                    docs = [doc for doc in st.session_state.documents if 'category' in doc]
                    labels = [doc['category'] for doc in docs]
                    
                    # Train model
                    st.session_state.classifier.train(docs, labels)
                    
                    # Success
                    st.success("Classifier trained successfully.")
                    
                    # Suggest saving
                    st.info("Don't forget to save your model in the sidebar.")
                except Exception as e:
                    st.error(f"Error training classifier: {str(e)}")

def create_settings_ui():
    """Create the settings UI."""
    st.header("Settings")
    
    # System settings
    st.subheader("System Settings")
    
    # Preprocessing settings
    st.markdown("### Text Preprocessing")
    use_stemming = st.checkbox("Use Stemming (instead of Lemmatization)", value=False)
    if use_stemming != getattr(st.session_state.indexer.preprocessor, 'use_stemming', False):
        st.session_state.indexer.preprocessor.use_stemming = use_stemming
        st.success("Preprocessing settings updated.")
    
    # Search settings
    st.markdown("### Search Settings")
    similarity_threshold = st.slider(
        "Similarity Threshold",
        min_value=0.0,
        max_value=1.0,
        value=st.session_state.search_engine.similarity_threshold,
        step=0.05
    )
    if similarity_threshold != st.session_state.search_engine.similarity_threshold:
        st.session_state.search_engine.similarity_threshold = similarity_threshold
        st.success("Search settings updated.")
    
    fuzzy_ratio_threshold = st.slider(
        "Fuzzy Match Threshold",
        min_value=0,
        max_value=100,
        value=st.session_state.search_engine.fuzzy_ratio_threshold,
        step=5
    )
    if fuzzy_ratio_threshold != st.session_state.search_engine.fuzzy_ratio_threshold:
        st.session_state.search_engine.fuzzy_ratio_threshold = fuzzy_ratio_threshold
        st.success("Fuzzy match settings updated.")
    
    # Advanced settings
    with st.expander("Advanced Settings"):
        # Vectorizer settings
        st.markdown("#### Vectorizer Settings")
        max_features = st.number_input(
            "Max Features",
            min_value=1000,
            max_value=50000,
            value=10000,
            step=1000
        )
        ngram_range = st.slider(
            "N-gram Range",
            min_value=1,
            max_value=3,
            value=(1, 2)
        )
        
        # Apply button
        if st.button("Apply Vectorizer Settings"):
            st.session_state.indexer.vectorizer = TfidfVectorizer(
                max_df=0.9,
                min_df=2,
                max_features=max_features,
                ngram_range=ngram_range
            )
            st.success("Vectorizer settings updated. You need to rebuild the index.")
        
        # Dimensionality reduction settings
        st.markdown("#### Dimensionality Reduction Settings")
        n_components = st.number_input(
            "Number of Components",
            min_value=50,
            max_value=500,
            value=100,
            step=10
        )
        
        # Apply button
        if st.button("Apply SVD Settings"):
            st.session_state.indexer.svd = TruncatedSVD(n_components=n_components)
            st.success("SVD settings updated. You need to rebuild the index.")
    
    # Reset all
    st.markdown("### Reset System")
    if st.button("Reset All"):
        st.session_state.indexer = DocumentIndexer()
        st.session_state.search_engine = FuzzySearchEngine(st.session_state.indexer)
        st.session_state.classifier = DocumentClassifier()
        st.session_state.documents = []
        st.success("System reset to initial state.")

# Run the app
if __name__ == "__main__":
    # Download NLTK resources
    download_nltk_resources()
    
    # Create and run the app
    create_streamlit_app()